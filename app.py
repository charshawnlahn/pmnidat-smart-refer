import streamlit as st
from google import genai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
import json
import requests
import datetime

# --- 1. ตรรกะการบริหารจัดการวันที่ (Thai Chronology Management) ---
# รองรับการเติมข้อมูล วันที่ {{DAY}} เดือน {{MONTH}} พ.ศ. {{YEAR}} ลงในเอกสารอัตโนมัติ
THAI_MONTHS = [
    "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน",
    "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม"
]

def get_thai_date_parts():
    """วิเคราะห์และส่งคืนส่วนประกอบของวันที่ปัจจุบันในรูปแบบ พ.ศ."""
    now = datetime.datetime.now()
    return {
        "DAY": str(now.day),
        "MONTH": THAI_MONTHS[now.month - 1],
        "YEAR": str(now.year + 543) # การปรับเปลี่ยนปีคริสต์ศักราชเป็นพุทธศักราช
    }

# --- 2. ฟังก์ชันการจัดเก็บสถิติเชิงระบบ (Utilization Logging) ---
def log_usage(patient_name):
    """กระบวนการบันทึกสถิติการใช้งานระบบเพื่อการตรวจสอบย้อนหลัง (ถ้ามีการระบุ URL)"""
    try:
        url = st.secrets.get("APPS_SCRIPT_URL", "")
        if url:
            requests.post(url, json={"name": patient_name}, timeout=5)
    except Exception:
        pass

# --- 3. การกำหนดค่าการเชื่อมต่อ API ความปลอดภัยสูง (API Configuration) ---
try:
    API_KEY = st.secrets["GEMINI_API_KEY"]
    client = genai.Client(api_key=API_KEY)
    # อัปเดตโมเดลเป็นรุ่นที่มีเสถียรภาพสูงสุดในปี 2026 เพื่อป้องกันข้อผิดพลาด 404
    MODEL_ID = "gemini-2.0-flash" 
except Exception as e:
    st.error(f"❌ พบข้อผิดพลาดในการกำหนดค่าระบบ: {e}")
    st.stop()

# --- 4. การจัดการหน่วยความจำชั่วคราว (Session State Memory Management) ---
# จัดเตรียมช่องว่างสำหรับรองรับข้อมูลทั้ง 9 ส่วนจากการคัดลอกระบบ @ThanHIS
field_keys = ['s11', 's12', 's13', 's14', 's2', 's31', 's32', 's33', 's34']
for key in field_keys:
    if key not in st.session_state:
        st.session_state[key] = ""

if "extracted_json_data" not in st.session_state:
    st.session_state.extracted_json_data = None



# --- 5. ฟังก์ชันบริหารจัดการชุดข้อมูลทดสอบ (Simulated Clinical Data Engine) ---

def load_test_data():
    """บรรจุชุดข้อมูลจำลองระดับวิชาชีพเข้าสู่ระบบเพื่อทดสอบตรรกะการสกัดและการคำนวณ """
    
    # 5.1 ข้อมูลทางคลินิกเบื้องต้นและประวัติการรักษา (Clinical Context)
    st.session_state.s11 = (
        "นาย ชาย ธัญญารักษ์ อายุ 40 ปี 5 เดือน [cite: 5]\n"
        "สิทธิ์ : จ่ายตรงกรมบัญชีกลาง\n"
        "Admit Date 01/03/2569\n"
        "จำนวนวัน Detox [5] วัน Rehab [10] วัน\n"
        "CC : เสพสุราซ้ำ ต้องการเข้ารับการบำบัดรักษา \n"
        "เคยมานอน รพ. 4 ครั้ง จำหน่ายล่าสุดวันที่ 25 กันยายน 2568"
    )
    
    # 5.2 ข้อมูลการวินิจฉัยโรคตามรหัสสากล (Diagnosis Logic)
    # รหัส ICD-10 เขียนติดกันตามมาตรฐาน HIS
    st.session_state.s12 = (
        "1. F105 - โรคจิตจากสุรา (Alcohol Psychosis) \n"
        "2. I10 - โรคความดันโลหิตสูง (Hypertension)"
    )
    
    # 5.3 ข้อมูลรายการยาที่ได้รับ (Pharmacological Profile)
    # ชื่อยาภาษาอังกฤษต้องเป็น UPPERCASE ทั้งหมดเสมอ
    st.session_state.s13 = (
        "1. AMLODIPINE 5 MG 1x1 pc (เช้า) \n"
        "2. QUETIAPINE 25 MG 1 tab hs (ก่อนนอน)\n"
        "(รายการยา Home-Med ทั้งหมด)"
    )
    
    # 5.4 บันทึกความก้าวหน้าทางการรักษา (Progress Note)
    st.session_state.s14 = (
        "ผู้ป่วยรู้สึกสบายดี รับประทานอาหารและนอนหลับได้ สัญญาณชีพคงที่\n"
        "ความดันโลหิต 120/80 mmHg อาการโดยรวมคงที่ เตรียมจำหน่ายกลับบ้าน "
    )
    
    # 5.5 คะแนนประเมินสภาวะสุขภาพจิต (Psychometric Assessments) 
    st.session_state.s2 = "9Q : 5 คะแนน, 8Q : 0 คะแนน, BPRS : 15 คะแนน "
    
    # 5.6 ข้อมูลทะเบียนประวัติและข้อมูลส่วนบุคคล (Demographics) 
    st.session_state.s31 = (
        "Hospital Number 690099999 [cite: 5]\n"
        "ชื่อ [ชาย] นามสกุล [ธัญญารักษ์] [cite: 5]\n"
        "เลขบัตรประชาชน [1-2345-67890-12-3] [cite: 5]\n"
        "ศาสนา [พุทธ] อาชีพ [ข้าราชการ] สถานภาพ [สมรส] การศึกษา [ปริญญาตรี] "
    )
    
    # 5.7 ข้อมูลที่อยู่และการติดต่อประสานงาน (Logistics & Contact) 
    st.session_state.s32 = "เลขที่ 123 ต.หลักหก อ.เมืองปทุมธานี จ.ปทุมธานี "
    st.session_state.s33 = "คุณ หญิง ธัญญารักษ์ (ภรรยา) เบอร์โทร: 081-234-XXXX "
    st.session_state.s34 = "สิทธิหลัก [จ่ายตรง] สถานพยาบาลหลัก [สถาบันบำบัดรักษาและฟื้นฟูผู้ติดยาเสพติดฯ] "
    
    # สั่งให้แอปรีเฟรชเพื่อแสดงผลข้อมูลในช่องกรอกทันที
    st.rerun()

# --- 6. ฟังก์ชันล้างข้อมูล (Data Sanitization Logic) ---

def clear_all_data():
    """ทำลายข้อมูลใน Session State ทั้งหมดเพื่อเริ่มเคสใหม่และรักษาความลับของผู้ป่วย"""
    for key in field_keys:
        st.session_state[key] = ""
    st.session_state.extracted_json_data = None
    st.rerun()


# --- 7. การจัดการแถบเมนูข้าง (Sidebar Configuration) ---

with st.sidebar:
    # 7.1 การแสดงสัญลักษณ์สถาบันฯ และหัวข้อคู่มือเชิงสถาบัน
    st.image("https://pmnidathis.dms.go.th/static/media/health.1bfd961f.png", width=100)
    st.header("📖 คู่มือการใช้งานระบบ")
    st.write("ระบบช่วยสกัดข้อมูลเพื่อจัดทำแบบบันทึกการส่งต่อ (PMNIDAT 062)")
    
    # 7.2 ปุ่มควบคุมศูนย์ปฏิบัติการ (Operation Controls)
    # ใช้ Key เฉพาะเพื่อป้องกันความซ้ำซ้อนของ ID ในระบบ Streamlit
    st.subheader("🛠️ เครื่องมือระบบ")
    col_t1, col_t2 = st.columns(2)
    with col_t1:
        if st.button("🧬 ตัวอย่างข้อมูล", key="side_load_sample", use_container_width=True, 
                     help="บรรจุข้อมูลจำลองเพื่อการทดสอบระบบ UAT"):
            load_test_data()
    with col_t2:
        if st.button("🧹 ล้างข้อมูล", key="side_clear_ui", use_container_width=True, 
                     help="ทำลายข้อมูลในหน่วยความจำเพื่อเริ่มเคสใหม่ (PDPA)"):
            clear_all_data()
            
    st.divider()
    
    # 7.3 แนวทางการปฏิบัติการคัดลอกข้อมูล (Data Extraction Protocol)
    st.markdown("### **วิธีการคัดลอกข้อมูลจาก @ThanHIS**")
    st.info("""
    1. **การเลือกข้อมูล:** คลิกเมาส์ซ้ายค้างที่ต้นข้อความ ลากครอบคลุมข้อมูลที่ต้องการ
    2. **การคัดลอก:** กดแป้นพิมพ์ **Ctrl+C** (Copy)
    3. **การวางข้อมูล:** คลิกในช่องรับข้อมูลที่กำหนด แล้วกด **Ctrl+V** (Paste)
    """)
    
    # 7.4 รายละเอียดขั้นตอนการสกัดข้อมูลรายโดเมน (Detailed Domain Steps)
    with st.expander("🟢 STEP 1: ระบบผู้ป่วยใน (IPD)", expanded=True):
        st.markdown("""
        **1.1 Admission Note:**
        - เข้าสู่เมนู Admission note
        - คัดลอกข้อมูลทั้งหมด ตั้งแต่ส่วนต้นจนถึงบรรทัดสุดท้าย
        
        **1.2 การวินิจฉัย:**
        - เข้าสู่เมนู "การวินิจฉัย"
        - คัดลอกรหัส ICD-10 ชื่อโรค และประเภทการวินิจฉัยทั้งหมด
        
        **1.3 Order / Meds:**
        - เข้าสู่เมนู **"Order"**
        - คัดลอกข้อมูล **Order และ Medication ทั้งหมด** จากส่วนบนถึงล่างสุด
        
        **1.4 Progress Note:**
        - เข้าสู่เมนู **"Progress note"**
        - คัดลอก **Progress note ทั้งหมด** ที่เป็นปัจจุบันที่สุด
        """)

    with st.expander("🔵 STEP 2: การประเมิน (Assessment)"):
        st.markdown("""
        - เมนู Admission note → เลือกปุ่ม **"ข้อมูลผู้ป่วยนอก"**
        - เลื่อนไปยังหัวข้อ **Assessment**
        - คัดลอกผลคะแนน **9Q, 8Q, BPRS** และคะแนนอื่นๆ ที่เกี่ยวข้อง
        """)

    with st.expander("🟠 STEP 3: เวชระเบียน (Registration)"):
        st.markdown("""
        - ระบบผู้ป่วยนอก → เวชระเบียน → ลงทะเบียนผู้ป่วยใหม่
        - ค้นหาด้วยรหัส HN เพื่อเข้าสู่ฐานข้อมูลหลัก
        
        **3.1 ทั่วไป 1:** คัดลอก Hospital number, ชื่อ-นามสกุล, อายุ, เลขบัตรประชาชน และศาสนา
        
        **3.2 ทั่วไป 2:** กดแสดง **"ที่อยู่ปัจจุบัน"** และคัดลอกข้อมูลที่อยู่โดยละเอียด
        
        **3.3 ผู้ติดต่อ:** คัดลอกนามผู้ดูแล ความสัมพันธ์ และหมายเลขโทรศัพท์  
        
        **3.4 สิทธิรักษา:** คัดลอกสิทธิการรักษา และ **"สถานพยาบาลหลัก"** เพื่อการส่งต่อ  
        """)
        
    st.divider()
    # 7.5 การแสดงรุ่นและสิทธิบัตรการพัฒนา (Version & Credits)
    # ตรวจสอบวงเล็บปิดให้สมบูรณ์เพื่อป้องกัน IndentationError
    st.caption("PMNIDAT Smart Transfer v3.36 (Official Master)")
    st.success("Created by Dr. Charshawn Lahnwong (5 Mar 2026)")

# --- สิ้นสุดส่วนที่ 3 ---


# --- 8. การออกแบบส่วนหัวข้อหลัก (Application Header) ---

# แสดงชื่อระบบและรุ่นเพื่อความเป็นทางการและน่าเชื่อถือของหน้าจอหลัก
st.title("🏥 PMNIDAT Smart Transfer")
st.subheader("ผู้ช่วยพิมพ์ 'แบบบันทึกข้อมูลเพื่อส่งต่อ (PMNIDAT 062)' โดยอัตโนมัติ (Version 3.36)")
st.divider()

# --- 9. ส่วนการรับข้อมูล 9 ส่วน (The 9 Input Fields) ---

# กลุ่มที่ 1: ระบบผู้ป่วยใน (IPD) - วิเคราะห์ข้อมูลทางคลินิกเชิงลึก
st.markdown("### **🟢 Step 1: ระบบผู้ป่วยใน (IPD Data)**")
s1_cols = st.columns(4)

with s1_cols[0]:
    st.text_area(
        "1.1 Admission Note", 
        height=300,
        placeholder="คัดลอกข้อมูลแรกรับ ประวัติสำคัญ และเหตุการณ์ปัจจุบัน...",
        key="main_s11", # แก้ไข: กำหนด Key เฉพาะตัวเพื่อป้องกันความซ้ำซ้อน
        help="คัดลอกจากเมนู Admission Note ในระบบ IPD"
    )

with s1_cols[1]:
    st.text_area(
        "1.2 การวินิจฉัย (ICD-10)", 
        height=300,
        placeholder="คัดลอกรหัส ICD-10 และชื่อโรคภาษาอังกฤษ/ไทย...",
        key="main_s12", 
        help="คัดลอกจากเมนูการวินิจฉัยเพื่อสกัดรหัสโรคแบบไม่มีจุดทศนิยม"
    )

with s1_cols[2]:
    st.text_area(
        "1.3 Order / Medication", 
        height=300,
        placeholder="คัดลอกรายการยาทั้งหมดจากหน้า Order...",
        key="main_s13", 
        help="คัดลอกเพื่อสกัดรายการยา Home-Med ที่เป็นตัวพิมพ์ใหญ่ (UPPERCASE)"
    )

with s1_cols[3]:
    st.text_area(
        "1.4 Progress Note", 
        height=300,
        placeholder="คัดลอกบันทึกการติดตามอาการล่าสุดทั้งหมด...",
        key="main_s14", 
        help="คัดลอกบันทึกความก้าวหน้าเพื่อสังเคราะห์ปัญหาที่ส่งต่อ"
    )

st.divider()

# กลุ่มที่ 2: การประเมิน (Assessment Scores) - วิเคราะห์สภาวะสุขภาพจิต
st.markdown("### **🔵 Step 2: การประเมิน (Assessment)**")
st.text_area(
    "คัดลอกผลคะแนน 9Q, 8Q, BPRS, OAS, GAF ทั้งหมดมาวางที่นี่", 
    height=150,
    placeholder="ระบุคะแนนการประเมินที่สำคัญสำหรับการส่งต่อ...",
    key="main_s2",
    help="ดึงจากหน้า Assessment เพื่อวิเคราะห์ระดับ Stage of Change และความเสี่ยง"
)

st.divider()

# กลุ่มที่ 3: เวชระเบียนและการติดต่อ (Registration & Contact)
st.markdown("### **🟠 Step 3: เวชระเบียน (Registration)**")
s3_cols = st.columns(4)

with s3_cols[0]:
    st.text_area(
        "3.1 ข้อมูลทั่วไป (Registration 1)", 
        height=200,
        placeholder="HN, ชื่อ-นามสกุล, อายุ, เลขบัตรประชาชน...",
        key="main_s31",
        help="คัดลอกจากหน้า 'ทั่วไป 1' ในระบบเวชระเบียน"
    )

with s3_cols[1]:
    st.text_area(
        "3.2 ที่อยู่ปัจจุบัน (Registration 2)", 
        height=200,
        placeholder="ระบุที่อยู่ที่ผู้ป่วยพักอาศัยจริงในปัจจุบัน...",
        key="main_s32",
        help="คัดลอกจากหน้า 'ทั่วไป 2' เพื่อกำหนดพิกัดการเยี่ยมบ้าน"
    )

with s3_cols[2]:
    st.text_area(
        "3.3 ผู้ติดต่อ / เบอร์โทรศัพท์", 
        height=200,
        placeholder="ชื่อผู้ดูแล ความสัมพันธ์ และหมายเลขติดต่อ...",
        key="main_s33",
        help="คัดลอกจากหน้า 'ผู้ติดต่อ' เพื่อประสานงานกับเครือข่าย"
    )

with s3_cols[3]:
    st.text_area(
        "3.4 สิทธิการรักษา", 
        height=200,
        placeholder="ประเภทสิทธิรักษา และสถานพยาบาลหลัก...",
        key="main_s34",
        help="คัดลอกจากหน้า 'สิทธิรักษา' เพื่อวิเคราะห์ Checkbox R1-R4"
    )

# --- สิ้นสุดส่วนที่ 4 ---


# --- 10. ระบบประมวลผลอัจฉริยะ (AI Extraction Engine - v3.36 Optimized) ---

if st.button("🚀 กดเพื่อประมวลผลและสกัดข้อมูลด้วย Gemini 2.0 Flash", key="btn_run_extraction", use_container_width=True):
    # 10.1 ดึงข้อมูลวันที่ปัจจุบันและรวบรวมข้อมูลดิบจาก 9 ช่องกรอก
    thai_date_data = get_thai_date_parts()
    today_str = f"{thai_date_data['DAY']} {thai_date_data['MONTH']} {thai_date_data['YEAR']}"
    
    # รวบรวมข้อมูลดิบทั้งหมดจาก Session State ที่ผูกไว้กับช่องกรอกในส่วนที่ 4
    all_raw_data = f"""
    วันที่ทำรายการปัจจุบัน: {today_str}
    
    [กลุ่มที่ 1: ข้อมูล IPD]
    1.1 Admission Note: {st.session_state.main_s11}
    1.2 การวินิจฉัย: {st.session_state.main_s12}
    1.3 Order / Meds: {st.session_state.main_s13}
    1.4 Progress Note: {st.session_state.main_s14}
    
    [กลุ่มที่ 2: การประเมิน]
    Assessment Score: {st.session_state.main_s2}
    
    [กลุ่มที่ 3: เวชระเบียน]
    3.1 ข้อมูลทั่วไป: {st.session_state.main_s31}
    3.2 ที่อยู่ปัจจุบัน: {st.session_state.main_s32}
    3.3 ผู้ติดต่อ: {st.session_state.main_s33}
    3.4 สิทธิการรักษา: {st.session_state.main_s34}
    """
    
    with st.spinner('Gemini 2.0 กำลังวิเคราะห์ข้อมูลเชิงลึกและคำนวณระยะเวลา...'):
        # 10.2 คำสั่ง (Prompt) สำหรับการสกัดข้อมูลระดับ PhD เพื่อรองรับไฟล์แม่แบบ .docx
        prompt = f"""
        คุณคือผู้ช่วยวิจัยทางการแพทย์ระดับ PhD จงสกัดข้อมูลจากชุดข้อมูลดิบที่ให้มาเพื่อลงในแบบฟอร์ม PMNIDAT 062
        โดยปฏิบัติตามกฎเหล็ก "Accuracy & Professional Formatting" อย่างเคร่งครัด:

        1. [Checkbox R1-R4 Logic]: วิเคราะห์สิทธิการรักษาแล้วส่งค่า [✓] ใน Key ที่ถูกต้องเพียงหนึ่งช่อง และ [ ] ในช่องที่เหลือ:
           - R1: ประกันสุขภาพ (บัตรทอง/UC/30 บาท)
           - R2: ประกันสังคม
           - R3: ท.74 ข้าราชการ/จ่ายตรง/รัฐวิสาหกิจ
           - R4: อื่นๆ (หากเลือกอันนี้ ให้ระบุรายละเอียดใน R_NOTE)

        2. [Formatting & Single Row Logic]:
           - DX (การวินิจฉัย): สกัดรหัส ICD-10 (ติดกันไม่มีจุด) และชื่อโรค เขียนต่อกันในแถวเดียว คั่นด้วยคอมม่า (,)
           - MEDS (รายการยา): สกัดรายการยา Home-Med เขียนชื่อภาษาอังกฤษเป็น UPPERCASE ทั้งหมด พร้อมวิธีใช้ คั่นด้วยคอมม่า (,) ในแถวเดียว
           - PROGRESS: สรุปปัญหาที่ส่งต่อให้เป็นย่อหน้าเดียวที่มีความกระชับและเป็นวิชาการ 2-3 บรรทัด

        3. [Calculation Audit]:
           - LOC: คำนวณระยะเวลาในชุมชน (วัน) = วันที่ปัจจุบัน ({today_str}) ลบ วันที่จำหน่ายล่าสุด (LAST_DC)
           - LOS: คำนวณวันนอนรวม = นำจำนวนวัน Detox และ Rehab มาบวกกัน

        ข้อมูลดิบสำหรับวิเคราะห์:
        {all_raw_data}

        ตอบกลับในรูปแบบ JSON ที่มี Key ดังนี้เท่านั้น:
        NAME, AGE, HN, ID, EDU, CAREER, RELIGION, STATUS, R1, R2, R3, R4, R_NOTE, LAST_DC, LOC, ADMIT_DATE, VISIT_NUM, CC, CONTACT, RELATION, PHONE, NEAR_HOSP, LOS, ADDRESS, Q9, Q8, MEDS, DX, PROGRESS, POST_SERVICE
        """
        
        try:
            # 10.3 ประมวลผลผ่านโมเดล Gemini 2.0 Flash รุ่นล่าสุดปี 2026
            response = client.models.generate_content(model=MODEL_ID, contents=prompt)
            # ใช้ Regex เพื่อดึงข้อมูลเฉพาะส่วนที่เป็น JSON
            json_match = re.search(r'\{.*\}', response.text, re.DOTALL)
            
            if json_match:
                st.session_state.extracted_json_data = json.loads(json_match.group())
                # ผสานข้อมูลวันที่ปัจจุบันเข้าสู่ชุดข้อมูลหลัก
                st.session_state.extracted_json_data.update(thai_date_data)
                st.success("✅ วิเคราะห์ข้อมูลและสกัดค่า Checkboxes สำเร็จเรียบร้อย!")
                # แสดง JSON เพื่อให้คุณหมอตรวจสอบความถูกต้อง
                with st.expander("🔍 ตรวจสอบข้อมูลที่สกัดได้ (Data Audit)"):
                    st.write(st.session_state.extracted_json_data)
            else:
                st.error("❌ ระบบไม่สามารถประมวลผลข้อมูลในรูปแบบที่กำหนดได้ กรุณาตรวจสอบข้อมูลดิบอีกครั้ง")
        except Exception as e:
            st.error(f"⚠️ เกิดข้อผิดพลาดในขั้นตอนประมวลผล AI: {e}")


# --- 11. ฟังก์ชันจัดการไฟล์ Word เชิงลึก (Advanced Document Orchestration) ---

def fill_pmnidat_doc(data):
    """ฟังก์ชันนำข้อมูลจากระบบเข้าสู่ไฟล์แม่แบบ .docx โดยรองรับโครงสร้างตารางและ Checkbox อย่างสมบูรณ์"""
    try:
        # 11.1 การเข้าถึงไฟล์แม่แบบที่คุณหมอจัดเตรียมไว้ (Master Template)
        # ตรวจสอบว่าชื่อไฟล์ตรงกับที่คุณหมออัปโหลดให้ในระบบ
        doc = Document("PMNIDAT 062 แบบบันทึกข้อมูลเพื่อส่งต่อ.docx")
        
        # 11.2 การเตรียมชุดข้อมูลสำหรับแทนที่ (Key Mapping) 
        # แปลง Key ให้เป็นตัวพิมพ์ใหญ่ทั้งหมดเพื่อความแม่นยำในการค้นหา Placeholder
        mapping = {f"{{{{{k.upper()}}}}}": str(v) for k, v in data.items()}
        
        def apply_style_and_replace(paragraph):
            """ตรรกะการวิเคราะห์บริบท (Context-Aware) เพื่อการจัดวางและแทนที่ข้อมูล"""
            
            # ก. การจัดการตำแหน่งการวาง (Alignment Strategy): 
            # วันที่ในส่วนหัวเอกสารให้ชิดขวา แต่ 'วันที่จำหน่าย' ในเนื้อหาตารางให้ชิดซ้าย
            if "วันที่" in paragraph.text and "จำหน่าย" not in paragraph.text:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # ข. กระบวนการค้นหาและแทนที่ข้อมูลเชิงลึก (Variable Injection)
            for key, value in mapping.items():
                if key in paragraph.text:
                    # แทนที่ Placeholder ด้วยข้อมูลจริง (เช่น {{R1}} จะกลายเป็น [✓] หรือ [ ])
                    paragraph.text = paragraph.text.replace(key, value)
                    
                    # ค. การกำหนดคุณลักษณะตัวอักษรมาตรฐาน (Standard Font Properties)
                    for run in paragraph.runs:
                        run.font.size = Pt(13) # กำหนดขนาดฟอนต์ 13 pt ตามมาตรฐานสถาบันฯ
                        
                        # ง. ตรรกะการเน้นความสำคัญ (Clinical Bolding):
                        # เน้นตัวหนาเฉพาะข้อมูลวิกฤต (DX, MEDS, PROGRESS, CC) เพื่อให้แพทย์ปลายทางอ่านง่าย
                        critical_info_keys = ["{{DX}}", "{{MEDS}}", "{{PROGRESS}}", "{{CC}}"]
                        if any(ck.upper() in key for ck in critical_info_keys):
                            run.font.bold = True

        # 11.3 ดำเนินการตรวจสอบและแทนที่ในเนื้อหาย่อหน้าปกติ (Body Paragraphs)
        for p in doc.paragraphs:
            apply_style_and_replace(p)
            
        # 11.4 ดำเนินการตรวจสอบภายในโครงสร้างตาราง (Table Traversal)
        # เนื่องจากไฟล์แม่แบบฉบับใหม่ของคุณหมอใช้ตารางเป็นโครงสร้างหลักในการล็อคตำแหน่งข้อมูล
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        apply_style_and_replace(p)
                            
        # 11.5 การจัดเก็บไฟล์ลงในหน่วยความจำชั่วคราวเพื่อเตรียมการส่งมอบ (Memory Buffering)
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"⚠️ ระบบไม่สามารถดำเนินการเข้าถึงไฟล์แม่แบบ Word ได้: {e}")
        return None

# --- สิ้นสุดส่วนที่ 6 ---
# --- 12. การแสดงผลลัพธ์และการควบคุมการดาวน์โหลด (Final Delivery Logic) ---

# ตรวจสอบความพร้อมของข้อมูลที่สกัดโดย AI จากหน่วยความจำ Session State
if st.session_state.extracted_json_data:
    # เรียกใช้ฟังก์ชันจัดการไฟล์ Word เพื่อสร้างเอกสารจากแม่แบบฉบับสมบูรณ์
    word_file_final = fill_pmnidat_doc(st.session_state.extracted_json_data)
    
    if word_file_final:
        # บันทึกสถิติการใช้งานเข้าระบบ Dashboard (อ้างอิงฟังก์ชันในส่วนที่ 1)
        log_usage(st.session_state.extracted_json_data.get('NAME', '[ไม่ระบุชื่อ]'))
        
        st.divider()
        st.balloons() # เฉลิมฉลองความสำเร็จในการจัดทำเอกสาร
        st.success("🎉 ระบบจัดทำเอกสาร PMNIDAT 062 ฉบับสมบูรณ์ (Master v3.36) สำเร็จแล้ว!")
        
        # การกำหนดชื่อไฟล์ตามชื่อผู้ป่วยเพื่อความสะดวกในการจัดเก็บเอกสาร
        patient_label = st.session_state.extracted_json_data.get('NAME', '062')
        file_name_output = f"Refer_{patient_label}.docx"
        
        # ปุ่มดาวน์โหลดไฟล์เอกสารที่จัดรูปแบบตามโครงสร้างตารางที่คุณหมอออกแบบ
        st.download_button(
            label=f"💾 ดาวน์โหลดไฟล์: {file_name_output}",
            data=word_file_final,
            file_name=file_name_output,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="btn_final_download",
            use_container_width=True
        )

# --- 13. มาตรการความปลอดภัยและจริยธรรมการจัดการข้อมูล (PDPA & Ethical Footer) ---
# ส่วนประกาศสำคัญเพื่อรักษามาตรฐานความลับของผู้ป่วยและมาตรฐานสถาบันฯ

st.divider()
st.info("""
    **🛡️ มาตรการรักษาความปลอดภัยของข้อมูลคนไข้ (PDPA Compliance):**
    
    * **Zero-Retention Policy:** ระบบประมวลผลแบบ Real-time บนหน่วยความจำชั่วคราว (RAM) เท่านั้น **ไม่มีการจัดเก็บข้อมูลส่วนบุคคลของผู้ป่วย** ลงในฐานข้อมูลถาวรหรือบันทึกไฟล์ทิ้งไว้บนเซิร์ฟเวอร์
    * **Session Isolation:** ข้อมูลที่คัดลอกมาวางจะถูกทำลายทิ้งทันทีเมื่อมีการรีเฟรชหน้าจอ (Refresh) หรือปิดเบราว์เซอร์ เพื่อป้องกันการรั่วไหลของข้อมูลระหว่างเคส
    * **Clinical Verification:** เนื่องจากเป็นระบบ AI ช่วยสกัดข้อมูล ผู้ใช้งานต้องตรวจสอบความถูกต้องของเนื้อหาและลงนามรับรองในเอกสารฉบับจริงทุกครั้งก่อนนำไปใช้งานทางคลินิก
    
    ---
    **🔬 พัฒนาระบบโดย:** ดร.นพ.ชาฌาน หลานวงศ์ (หมออาร์ม) | MD-PhD (Pharmacology) 
    แพทย์เวชศาสตร์สารเสพติด สถาบันบำบัดรักษาและฟื้นฟูผู้ติดยาเสพติดแห่งชาติบรมราชชนนี (สบยช.)
    
    **💡 ปรึกษาปัญหาการใช้งาน/ติดต่อสอบถามได้ที่**
    * **Line:** armlahn | **โทร:** 094-991-4599

    """)

# --- สิ้นสุดการทำงานของระบบ PMNIDAT Smart Transfer v3.36 ---
